#!/usr/bin/env python3
"""
SOY ANALISTA — Newsletter DOCX Generator
Argentina Mundial 2026 — Type 1: Informe de Seleccion
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── BRAND COLORS ──
NEGRO = RGBColor(0x1C, 0x19, 0x15)
CAFE = RGBColor(0x6B, 0x5B, 0x4D)
CAFE_OSCURO = RGBColor(0x4A, 0x3E, 0x34)
TIERRA = RGBColor(0xA6, 0x98, 0x82)
ROJO = RGBColor(0xC1, 0x36, 0x28)
VERDE = RGBColor(0x3A, 0x7D, 0x44)
DORADO = RGBColor(0xB8, 0x92, 0x2F)
AZUL = RGBColor(0x2E, 0x5E, 0x8C)
BLANCO = RGBColor(0xFE, 0xFC, 0xF8)
CREMA_BG = "F5F0E8"
CREMA_OSCURA_BG = "EDE6D8"
ARENA_BG = "D9CEBC"
NEGRO_BG = "1C1915"


def pct_color(pct):
    if pct >= 75: return ROJO
    if pct >= 50: return DORADO
    if pct >= 25: return AZUL
    return TIERRA


def set_cell_bg(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1, color=NEGRO, size=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
    return h


def add_para(doc, text, bold=False, italic=False, color=NEGRO, size=10, align=None, space_after=6):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p


def add_mixed_para(doc, parts, size=10, align=None, space_after=6):
    """parts = list of (text, bold, italic, color)"""
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic, color in parts:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic
    return p


def add_editorial_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("━" * 80)
    run.font.size = Pt(4)
    run.font.color.rgb = NEGRO
    return p


def add_diamond_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─────────────── ◆ ───────────────")
    run.font.size = Pt(8)
    run.font.color.rgb = TIERRA
    return p


def add_section_header(doc, number, title):
    add_editorial_rule(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run_num = p.add_run(f"{number}  ")
    run_num.font.size = Pt(9)
    run_num.font.color.rgb = TIERRA
    run_title = p.add_run(title)
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = NEGRO
    run_title.bold = True


def add_percentile_row(table, label, value, pct):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = str(value)
    row.cells[2].text = f"P{pct}"
    bar_chars = int(pct / 5)
    row.cells[3].text = "█" * bar_chars + "░" * (20 - bar_chars)
    # Style
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.font.color.rgb = CAFE
                elif i == 1:
                    run.font.color.rgb = NEGRO
                    run.bold = True
                elif i == 2:
                    run.font.color.rgb = pct_color(pct)
                    run.bold = True
                elif i == 3:
                    run.font.color.rgb = pct_color(pct)
                    run.font.size = Pt(7)


def generate_newsletter():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Georgia'
    style.font.size = Pt(10)
    style.font.color.rgb = NEGRO

    # ══════════════════════════════════════════════════════════════
    # HEADER
    # ══════════════════════════════════════════════════════════════

    # Brand header table (simulates dark header)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in header_table.rows[0].cells:
        set_cell_bg(cell, NEGRO_BG)
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(8)

    cell_left = header_table.rows[0].cells[0]
    p = cell_left.paragraphs[0]
    run = p.add_run("[A] Soy Analista")
    run.font.size = Pt(12)
    run.font.color.rgb = BLANCO
    run.bold = True

    cell_right = header_table.rows[0].cells[1]
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Informe de seleccion · Mundial 2026 · Tier 1")
    run.font.size = Pt(8)
    run.font.color.rgb = TIERRA
    run.italic = True

    # Team name block
    team_table = doc.add_table(rows=4, cols=1)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in team_table.rows:
        set_cell_bg(row.cells[0], NEGRO_BG)
        for para in row.cells[0].paragraphs:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)

    # Flag
    p = team_table.rows[0].cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("🇦🇷")
    run.font.size = Pt(28)

    # Team name
    p = team_table.rows[1].cells[0].paragraphs[0]
    run = p.add_run("Argentina")
    run.font.size = Pt(36)
    run.font.color.rgb = BLANCO
    run.bold = True

    # Subtitle
    p = team_table.rows[2].cells[0].paragraphs[0]
    run = p.add_run("CONMEBOL · FIFA #2 · Grupo J · Campeona defensora")
    run.font.size = Pt(11)
    run.font.color.rgb = TIERRA
    run.italic = True

    # Stats row
    p = team_table.rows[3].cells[0].paragraphs[0]
    p.paragraph_format.space_after = Pt(12)
    stats = [("4-3-3", "Sistema"), ("1ro", "Elim."), ("31GF-10GC", "18 PJ"), ("70.6%", "Win rate")]
    for val, label in stats:
        run = p.add_run(f"  {val} ")
        run.font.size = Pt(12)
        run.font.color.rgb = BLANCO
        run.bold = True
        run = p.add_run(f"({label})  ")
        run.font.size = Pt(8)
        run.font.color.rgb = TIERRA
        run.italic = True

    # Edition line
    add_para(doc, "Edicion #001 · 4 de marzo de 2026 · soyanalista.com",
             italic=True, color=TIERRA, size=9, align='center', space_after=12)

    # ══════════════════════════════════════════════════════════════
    # FREE SECTION
    # ══════════════════════════════════════════════════════════════

    add_section_header(doc, "02", "Resumen ejecutivo")

    add_para(doc, "Argentina de Lionel Scaloni llega al Mundial 2026 como campeona defensora y maxima candidata al titulo. Opera en un 4-3-3 flexible que muta al 4-4-2 sin balon y al 4-2-3-1 cuando busca control posicional. Con un 65-68% de posesion promedio (P88 entre selecciones mundialistas), domina los partidos desde la construccion pausada pero con verticalidad letal cuando detecta espacios a la espalda de la linea rival.")

    add_mixed_para(doc, [
        ("Como ataca: ", True, False, NEGRO),
        ("La profundidad ofensiva es la principal virtud: ", False, False, NEGRO),
        ("Messi (ICI 3.85) ", True, False, ROJO),
        ("sigue siendo el catalizador creativo a los 38 anos, pero ", False, False, NEGRO),
        ("Alvarez (ICI 3.50) ", True, False, NEGRO),
        ("y ", False, False, NEGRO),
        ("Lautaro Martinez (ICI 3.52) ", True, False, NEGRO),
        ("diversifican la finalizacion. Generan ", False, False, NEGRO),
        ("1.72 goles/partido ", True, False, NEGRO),
        ("en eliminatorias con 31 goles entre al menos ", False, False, NEGRO),
        ("12 anotadores distintos.", True, False, NEGRO),
    ])

    add_mixed_para(doc, [
        ("Como defiende: ", True, False, NEGRO),
        ("La mejor defensa de CONMEBOL: ", False, False, NEGRO),
        ("0.56 GA/partido (P91) ", True, False, ROJO),
        ("con ", False, False, NEGRO),
        ("12 porterias a cero ", True, False, NEGRO),
        ("en 18 partidos (66.7%). Romero y Lisandro Martinez forman una dupla de centrales de nivel Champions. El mediocampo Enzo-De Paul filtra progresiones rivales con la combinacion mas completa del torneo.", False, False, NEGRO),
    ])

    # Claves callout
    callout_table = doc.add_table(rows=1, cols=1)
    set_cell_bg(callout_table.rows[0].cells[0], NEGRO_BG)
    cell = callout_table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Claves del partido:")
    run.font.size = Pt(10)
    run.font.color.rgb = BLANCO
    run.bold = True
    run.italic = True

    claves = [
        "◆ Explotar la transicion defensiva cuando pierde en campo rival",
        "◆ Neutralizar a Messi alejandolo del centro",
        "◆ Evitar duelos aereos donde Romero domina",
        "◆ El plan B en el mediocampo es significativamente inferior al titular",
    ]
    for clave in claves:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(clave)
        run.font.size = Pt(9)
        run.font.color.rgb = BLANCO

    # ── Stat cards ──
    add_diamond_separator(doc)

    h = doc.add_heading("Los numeros en un vistazo", level=2)
    for run in h.runs:
        run.font.color.rgb = NEGRO

    stat_table = doc.add_table(rows=2, cols=4)
    stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [("38", "Puntos elim."), ("12-2-4", "V - E - D"), ("+21", "Dif. goles"), ("1ro", "Pos. CONMEBOL")]
    for i, (val, label) in enumerate(values):
        cell = stat_table.rows[0].cells[i]
        set_cell_bg(cell, CREMA_OSCURA_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(val)
        run.font.size = Pt(18)
        run.font.color.rgb = ROJO if i in [0, 3] else NEGRO
        run.bold = True

        cell2 = stat_table.rows[1].cells[i]
        set_cell_bg(cell2, CREMA_OSCURA_BG)
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(8)
        run2 = p2.add_run(label)
        run2.font.size = Pt(8)
        run2.font.color.rgb = TIERRA
        run2.italic = True

    # Metric summary
    metric_table = doc.add_table(rows=1, cols=1)
    set_cell_bg(metric_table.rows[0].cells[0], CREMA_BG)
    p = metric_table.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("0.56 GA/p (P91) · 66% posesion (P88) · 1.72 G/p (P82) · 12 porterias a cero")
    run.font.size = Pt(9)
    run.font.color.rgb = CAFE_OSCURO

    # ── Group ──
    add_diamond_separator(doc)
    h = doc.add_heading("El grupo", level=2)
    for run in h.runs: run.font.color.rgb = NEGRO

    add_mixed_para(doc, [
        ("Grupo J", True, False, NEGRO),
        (": Argentina 🇦🇷 · Algeria 🇩🇿 · Austria 🇦🇹 · Jordan 🇯🇴", False, False, NEGRO),
    ])

    group_table = doc.add_table(rows=5, cols=3)
    group_table.style = 'Light Shading'
    headers = ["Seleccion", "FIFA", "Perfil"]
    for i, h_text in enumerate(headers):
        group_table.rows[0].cells[i].text = h_text
        for run in group_table.rows[0].cells[i].paragraphs[0].runs:
            run.font.size = Pt(9)
            run.bold = True

    group_data = [
        ("🇦🇷 Argentina", "#2", "Posesion + pressing selectivo"),
        ("🇩🇿 Algeria", "#34", "Bloque medio, transiciones rapidas"),
        ("🇦🇹 Austria", "#25", "Pressing alto Rangnick, intensidad"),
        ("🇯🇴 Jordan", "#62", "Bloque bajo, debutante mundialista"),
    ]
    for r, (team, rank, profile) in enumerate(group_data, 1):
        group_table.rows[r].cells[0].text = team
        group_table.rows[r].cells[1].text = rank
        group_table.rows[r].cells[2].text = profile
        for cell in group_table.rows[r].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    add_mixed_para(doc, [
        ("Partido clave: ", True, False, NEGRO),
        ("Argentina vs Austria (J2, 21 jun, Arlington). Pressing Rangnick vs salida de balon argentina.", False, False, NEGRO),
    ])
    add_mixed_para(doc, [
        ("Probabilidad de clasificacion: ", True, False, NEGRO),
        ("1ro del grupo (82%)", True, False, VERDE),
    ])

    # ── CTA ──
    doc.add_paragraph()
    cta_table = doc.add_table(rows=1, cols=1)
    set_cell_bg(cta_table.rows[0].cells[0], CREMA_BG)
    cell = cta_table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("El informe completo incluye: dashboard estadistico con percentiles vs las 47 selecciones, analisis de 6 fases tacticas, fichas de 14 jugadores clave con ICI compuesto, 4 fortalezas y 4 debilidades, y XI probable con recomendaciones.")
    run.font.size = Pt(9)
    run.font.color.rgb = CAFE
    run.italic = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run("Son 18 paginas de analisis profesional — el nivel de un departamento de scouting.")
    run2.font.size = Pt(9)
    run2.font.color.rgb = CAFE
    run2.bold = True
    run2.italic = True

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(12)
    run3 = p3.add_run("[ Suscribite por 8 EUR/mes → ]")
    run3.font.size = Pt(12)
    run3.font.color.rgb = ROJO
    run3.bold = True

    # ── PAYWALL ──
    doc.add_paragraph()
    pw_table = doc.add_table(rows=1, cols=1)
    set_cell_bg(pw_table.rows[0].cells[0], NEGRO_BG)
    p = pw_table.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("🔒 El resto de este informe es exclusivo para suscriptores.")
    run.font.size = Pt(10)
    run.font.color.rgb = BLANCO

    # ══════════════════════════════════════════════════════════════
    # PAID SECTION — DASHBOARD
    # ══════════════════════════════════════════════════════════════

    add_section_header(doc, "07", "Dashboard estadistico")
    add_para(doc, "Percentiles calculados contra las 47 selecciones clasificadas al Mundial 2026.",
             italic=True, color=TIERRA, size=8)

    # Defense percentiles
    for section_title, metrics in [
        ("Fase defensiva", [
            ("Goles enc. / partido", "0.56", 91),
            ("Porterias a cero", "66.7%", 93),
            ("GA local / partido", "0.22", 96),
            ("GA visitante / partido", "0.89", 65),
            ("Duelos def. ganados", "~54%", 68),
        ]),
        ("Distribucion", [
            ("Posesion promedio", "~66%", 88),
            ("Precision de pase", "~87%", 82),
            ("Pases progresivos / p", "~8.5", 78),
            ("Pases al ult. tercio / p", "~12", 80),
        ]),
        ("Fase ofensiva", [
            ("Goles / partido", "1.72", 82),
            ("xG / partido (est.)", "~1.65", 80),
            ("Tiros / partido (est.)", "~13.5", 72),
            ("Conversion de gol", "~12.7%", 78),
            ("Goleadores distintos", "12+", 92),
        ]),
    ]:
        h = doc.add_heading(section_title, level=3)
        for run in h.runs: run.font.color.rgb = CAFE_OSCURO

        pct_table = doc.add_table(rows=1, cols=4)
        pct_table.rows[0].cells[0].text = "Metrica"
        pct_table.rows[0].cells[1].text = "Valor"
        pct_table.rows[0].cells[2].text = "Pct"
        pct_table.rows[0].cells[3].text = "Barra"
        for cell in pct_table.rows[0].cells:
            set_cell_bg(cell, CREMA_OSCURA_BG)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = TIERRA
                    run.bold = True

        for label, value, pct in metrics:
            add_percentile_row(pct_table, label, value, pct)

    # Dashboard reading
    callout2 = doc.add_table(rows=1, cols=1)
    set_cell_bg(callout2.rows[0].cells[0], NEGRO_BG)
    p = callout2.rows[0].cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Lectura global: Perfil top-5 mundial en todas las fases. Solidez defensiva (P91), posesion dominante (P88), produccion diversificada (P92). Inconsistencia: disparidad local/visitante en GA.")
    run.font.size = Pt(9)
    run.font.color.rgb = BLANCO
    run.italic = True
    run.bold = True

    # ══════════════════════════════════════════════════════════════
    # MODELO DE JUEGO
    # ══════════════════════════════════════════════════════════════

    add_section_header(doc, "08-12", "Modelo de juego")

    for subtitle, text in [
        ("Construccion", "Argentina construye desde atras con Romero y Lisandro abiertos a 25-30m. Enzo o De Paul baja entre centrales para trivote de salida (3v2). Laterales proyectan alto. Posesion ~66% (P88) con dos velocidades: fase paciente (70%, circulacion lateral) y fase de activacion (30%, pase vertical via Messi, Enzo o Mac Allister). E. Martinez participa con pases largos diagonales."),
        ("Progresion", "Opera por el centro via triangulo Enzo-De Paul-Mac Allister. Enzo orquesta (8-9 pases progresivos/p, 8G PL 25-26). De Paul conecta (motor, +11 km/p). Mac Allister alterna interior-mediapunta. Vias: canal central (40%), canal derecho (35%, Molina + Messi), canal izquierdo (25%, Nico Gonzalez + Tagliafico)."),
        ("Finalizacion", "~13.5 tiros/partido (P72), 1.72 goles/partido (P82). Perfil eficiente: calidad de finalizadores elite. xG ~1.65/p — ligero sobrerendimiento (+0.07). Tipos de gol: combinado central (~40%), jugada individual (~20%), transicion (~15%), balon parado (~15%), centro + remate (~10%)."),
        ("Pressing y defensa", "Bloque medio-alto. PPDA est. ~10-12 (moderado-alto). Triggers: pase hacia atras, recepcion de espaldas, lateral en banda, saque de puerta corto. Romero (agresivo) + Lisandro (posicional) = dupla complementaria. Debilidad: 4 derrotas (todas visitante) incluyeron goles en transiciones. Espacio a espalda de Molina = via de ataque."),
        ("Transiciones", "Ofensiva: letal tras recuperacion alta — pase filtrado a Messi/Alvarez en 3-4 seg. Defensiva (punto mas vulnerable): si pressing post-perdida falla, laterales proyectados dejan espacio. Las 4 derrotas coinciden con exposicion en ese espacio."),
    ]:
        h = doc.add_heading(subtitle, level=3)
        for run in h.runs: run.font.color.rgb = CAFE_OSCURO
        add_para(doc, text)

    # ══════════════════════════════════════════════════════════════
    # JUGADORES CLAVE
    # ══════════════════════════════════════════════════════════════

    add_section_header(doc, "14-15", "Jugadores clave")

    players = [
        ("👑 Lionel Messi", "RW/SS", "Inter Miami", 38, 3.85, "ELITE",
         "29G+19A MLS 2025 · 8G eliminatorias",
         [("Creativo", 3.9), ("Finaliz.", 3.8), ("Progres.", 3.0), ("Defensivo", 0.8)],
         "Capitan. A los 38 no presiona ni corre como antes pero sigue siendo el jugador mas peligroso del torneo con balon en el ultimo tercio. El carril derecho con Molina + Messi cortando adentro es la via de ataque principal.",
         "No dejarle recibir de frente entre lineas. Doblar marca. Alejarlo del centro hacia la banda."),
        ("⚡ Julian Alvarez", "ST", "Atletico Madrid", 26, 3.50, "ELITE",
         "13G+5A all comps 25-26 · 5G eliminatorias",
         [("Finaliz.", 3.5), ("Progres.", 3.2), ("Creativo", 2.8), ("Defensivo", 2.8)],
         "Delantero mas completo del plantel. Remata de todas posiciones, presiona centrales, conecta bajando al mediocampo. Nunca esta quieto.",
         "No marcarlo al hombre — arrastra al central. Mantener la linea."),
        ("🎯 Enzo Fernandez", "CM", "Chelsea (PL)", 25, 3.40, "ALTO",
         "8G+2A PL 25-26",
         [("Creativo", 3.8), ("Progres.", 3.5), ("Defensivo", 3.2), ("Finaliz.", 2.4)],
         "Corazon del equipo. Orquestador: recibe de espaldas, gira, distribuye. Lidera pases progresivos. Llegada al area con remate de media distancia.",
         "Presionarlo cuando recibe de espaldas. No darle tiempo para girar."),
        ("🔒 Cristian Romero", "CB", "Tottenham (PL)", 27, 3.10, "ALTO",
         "Cap. Spurs · 92% pases",
         [("Defensivo", 3.8), ("Progres.", 2.9), ("Creativo", 0.8), ("Finaliz.", 0.4)],
         "Central agresivo de anticipacion. Conduce desde atras. Duelo aereo dominante. 2 rojas PL 25-26 — estilo con riesgo pero otorga linea mas adelantada.",
         "Jugadores rapidos en diagonal explotan sus salidas. Vulnerable a dribblers."),
    ]

    for name, pos, club, age, ici, ici_label, stats_line, roles, profile, rival in players:
        # Player header
        p_card = doc.add_table(rows=1, cols=2)
        set_cell_bg(p_card.rows[0].cells[0], CREMA_OSCURA_BG)
        set_cell_bg(p_card.rows[0].cells[1], CREMA_OSCURA_BG)

        p = p_card.rows[0].cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(name)
        run.font.size = Pt(14)
        run.font.color.rgb = NEGRO
        run.bold = True
        p2 = p_card.rows[0].cells[0].add_paragraph()
        run2 = p2.add_run(f"{pos} · {club} · {age} anos")
        run2.font.size = Pt(8)
        run2.font.color.rgb = TIERRA
        run2.italic = True

        p = p_card.rows[0].cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f"{ici}")
        run.font.size = Pt(18)
        run.font.color.rgb = ROJO
        run.bold = True
        p3 = p_card.rows[0].cells[1].add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run3 = p3.add_run(f"ICI {ici_label}")
        run3.font.size = Pt(7)
        run3.font.color.rgb = TIERRA

        # Stats line
        add_para(doc, stats_line, italic=True, color=TIERRA, size=8, space_after=4)

        # Roles
        role_parts = []
        for role_name, role_val in roles:
            bar = "█" * int(role_val * 2.5) + "░" * (10 - int(role_val * 2.5))
            role_parts.append((f"{role_name}: {role_val}/4 {bar}  ", False, False, pct_color(int(role_val * 25))))
        add_mixed_para(doc, role_parts, size=8, space_after=4)

        # Profile
        add_para(doc, profile, size=9.5)

        # Rival key
        rival_parts = [
            ("Para el rival: ", True, False, ROJO),
            (rival, False, False, CAFE),
        ]
        add_mixed_para(doc, rival_parts, size=9)

        add_diamond_separator(doc)

    # Suplentes
    add_para(doc, "Tambien clave: Lautaro Martinez (ICI 3.52, ST, 14G Serie A), E. Martinez (ICI 3.22, GK, especialista penales), Mac Allister (ICI 3.18, CM, 3G UCL), De Paul (ICI 2.92, CM, motor y cohesion).",
             italic=True, color=TIERRA, size=9)

    # ══════════════════════════════════════════════════════════════
    # FORTALEZAS Y DEBILIDADES
    # ══════════════════════════════════════════════════════════════

    add_section_header(doc, "16", "Fortalezas y debilidades")

    add_para(doc, "✅ FORTALEZAS — QUE RESPETAR", bold=True, color=VERDE, size=11)

    sw_data = [
        (True, "1. Madurez competitiva incomparable",
         "Nucleo de 8-9 jugadores con 50+ partidos juntos. 4 titulos 2021-2024. No hay contexto que los desestabilice.",
         "70.6% win rate en ~85 partidos con Scaloni"),
        (True, "2. Profundidad ofensiva elite",
         "Messi, Alvarez, Lautaro: tres delanteros top-20. Garnacho, Almada, Nico Paz desde el banco. 12 goleadores distintos.",
         "1.72 goles/partido (P82), 12+ goleadores (P92)"),
        (True, "3. Solidez defensiva de elite",
         "Romero-Lisandro de nivel Champions. E. Martinez, seguridad y penales. Mediocampo Enzo-De Paul filtra progresiones.",
         "0.56 GA/p (P91). 0.22 GA/p local (P96). 66.7% arco en cero"),
        (True, "4. Pressing coordinado e inteligente",
         "Elige cuando apretar con triggers claros. Recuperaciones en campo rival alimentan transiciones ofensivas letales.",
         "Pressing selectivo — elige el momento, no se desgasta"),
    ]

    for is_str, title, desc, metric in sw_data:
        color = VERDE
        add_para(doc, title, bold=True, color=NEGRO, size=11, space_after=2)
        add_para(doc, desc, size=9.5, space_after=2)
        # Metric in colored box
        m_table = doc.add_table(rows=1, cols=1)
        set_cell_bg(m_table.rows[0].cells[0], CREMA_BG)
        p = m_table.rows[0].cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(metric)
        run.font.size = Pt(8)
        run.font.color.rgb = color
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_para(doc, "⚠️ DEBILIDADES — QUE EXPLOTAR", bold=True, color=ROJO, size=11)

    weak_data = [
        ("1. Vulnerabilidad de visitante",
         "8W-1E-0L en casa vs 4W-1E-4L de visitante. Las 4 derrotas fueron todas fuera.",
         "Win rate local 88.9% vs visitante 44.4%"),
        ("2. Transiciones defensivas expuestas",
         "Cuando pierde en campo rival con laterales proyectados, el espacio entre lineas se abre.",
         "Espalda de Molina (RB) = principal via de ataque"),
        ("3. Dependencia del mediocampo titular",
         "Caida significativa sin Enzo y De Paul. Paredes ofrece control pero no la misma dinamica.",
         "~80% win rate con ambos vs ~55% sin alguno"),
        ("4. Gestion fisica de Messi",
         "A los 38, Scaloni gestiona su carga. En un Mundial cada 3-4 dias, la dosificacion abre ventana.",
         "Messi jugo 12 de 18 partidos de eliminatorias"),
    ]

    for title, desc, metric in weak_data:
        add_para(doc, title, bold=True, color=NEGRO, size=11, space_after=2)
        add_para(doc, desc, size=9.5, space_after=2)
        m_table = doc.add_table(rows=1, cols=1)
        set_cell_bg(m_table.rows[0].cells[0], CREMA_BG)
        p = m_table.rows[0].cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(metric)
        run.font.size = Pt(8)
        run.font.color.rgb = ROJO
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ══════════════════════════════════════════════════════════════
    # XI PROBABLE
    # ══════════════════════════════════════════════════════════════

    add_section_header(doc, "17", "XI probable")

    # Lineup in dark box
    xi_table = doc.add_table(rows=5, cols=1)
    for row in xi_table.rows:
        set_cell_bg(row.cells[0], NEGRO_BG)

    p = xi_table.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("4-3-3 — Formacion mas utilizada (est. 14 de 18 eliminatorias)")
    run.font.size = Pt(9)
    run.font.color.rgb = TIERRA

    lines = [
        ("ATAQUE", "N. Gonzalez    J. Alvarez    L. Messi (C)"),
        ("MEDIOCAMPO", "A. Mac Allister    R. De Paul    Enzo Fernandez"),
        ("DEFENSA", "Tagliafico    L. Martinez    Romero    Molina"),
        ("PORTERO", "E. Martinez"),
    ]
    for i, (label, players_line) in enumerate(lines):
        p = xi_table.rows[i+1].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(label + "\n")
        run.font.size = Pt(7)
        run.font.color.rgb = TIERRA
        run2 = p.add_run(players_line)
        run2.font.size = Pt(11)
        run2.font.color.rgb = BLANCO
        run2.bold = True

    add_mixed_para(doc, [
        ("Alternativas: ", True, False, NEGRO),
        ("4-2-3-1 con Paredes (vs pressing alto). Doble 9 Alvarez + Lautaro (buscar gol).", False, False, NEGRO),
    ])

    # ── CLAVES TABLE ──
    add_diamond_separator(doc)
    add_para(doc, "CLAVES PARA ENFRENTAR A ARGENTINA", bold=True, color=ROJO, size=11)

    claves_table = doc.add_table(rows=7, cols=2)
    claves_data = [
        ("Presionar alto la salida", "Centrales dudan con presion, pierden balones peligrosos"),
        ("Atacar espalda de Molina", "Sube constantemente — espacio a su espalda explotable"),
        ("No seguir a Alvarez", "Mantener la linea; seguirlo desordena la defensa"),
        ("Anular a Enzo", "Sin Enzo, Argentina pierde su orquestador principal"),
        ("Messi: alejarlo del centro", "Empujarlo a la banda reduce angulo de pase y tiro"),
        ("Evitar penales", "E. Martinez en tandas es ventaja Argentina"),
        ("Exigir fisicamente", "Messi a los 38 se desgasta con intensidad 90 min"),
    ]
    for r, (clave, expl) in enumerate(claves_data):
        claves_table.rows[r].cells[0].text = clave
        claves_table.rows[r].cells[1].text = expl
        for para in claves_table.rows[r].cells[0].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                run.bold = True
                run.font.color.rgb = NEGRO
        for para in claves_table.rows[r].cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = CAFE
                run.italic = True

    # ── Download CTA ──
    doc.add_paragraph()
    dl_table = doc.add_table(rows=1, cols=1)
    set_cell_bg(dl_table.rows[0].cells[0], CREMA_BG)
    p = dl_table.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("[ Descargar informe completo — 18 pag. PDF ↓ ]")
    run.font.size = Pt(13)
    run.font.color.rgb = NEGRO
    run.bold = True
    p2 = dl_table.rows[0].cells[0].add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run("Incluye las 17 secciones, dashboard, fichas ICI y recomendaciones tacticas.")
    run2.font.size = Pt(8)
    run2.font.color.rgb = TIERRA
    run2.italic = True

    # ── Share ──
    share_table = doc.add_table(rows=1, cols=1)
    set_cell_bg(share_table.rows[0].cells[0], NEGRO_BG)
    p = share_table.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Si este analisis te parecio util, compartilo con alguien que necesite entender el Mundial con datos.")
    run.font.size = Pt(10)
    run.font.color.rgb = BLANCO
    run.italic = True
    run.bold = True

    add_para(doc, "¿Que seleccion queres que analicemos la semana que viene? Responde a este email con tu pedido.",
             italic=True, color=TIERRA, size=9, space_after=16)

    # ── Footer ──
    add_editorial_rule(doc)
    add_para(doc, "Fuentes: FBref, Transfermarkt, FIFA, CONMEBOL. Datos actualizados al 3 de marzo de 2026.",
             italic=True, color=TIERRA, size=8, align='center')
    add_para(doc, "Los percentiles son estimaciones basadas en datos publicos. ICI compuestos: 40% seleccion + 60% club.",
             italic=True, color=TIERRA, size=8, align='center')
    add_para(doc, "Soy Analista · soyanalista.com · @analistasoy",
             bold=True, color=NEGRO, size=10, align='center')

    # ── Save ──
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "newsletter_argentina_mundial2026.docx")
    doc.save(output_path)
    print(f"DOCX generado: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_newsletter()
